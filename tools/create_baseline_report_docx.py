from __future__ import annotations

import argparse
import csv
import json
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_BASELINE_DIR = PROJECT_ROOT / "results" / "analysis" / "baselines"
DEFAULT_RESULTS_DIR = PROJECT_ROOT / "results" / "final"
DEFAULT_FIGURES_DIR = PROJECT_ROOT / "figures" / "baselines" / "all_methods"
DEFAULT_REPORT_PATH = PROJECT_ROOT / "docs" / "reports" / "冻结基线实验设计与结果分析报告.docx"


METHOD_ORDER = ["dense", "file_fts", "graph_path", "fusion", "oracle"]
METHOD_LABELS = {
    "dense": "Dense / MiniLM",
    "file_fts": "File-FTS",
    "graph_path": "Graph-Path",
    "fusion": "Fusion / RRF",
    "oracle": "Oracle / 理想上限",
}
TASK_ORDER = ["semantic_fact", "multi_hop_relation", "exact_file_lookup"]
TASK_LABELS = {
    "semantic_fact": "Semantic Fact",
    "multi_hop_relation": "Multi-hop Relation",
    "exact_file_lookup": "Exact/File Lookup",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as file:
        return [dict(row) for row in csv.DictReader(file)]


def read_jsonl(path: Path) -> Iterable[dict[str, Any]]:
    with path.open("r", encoding="utf-8") as file:
        for line in file:
            line = line.strip()
            if line:
                yield json.loads(line)


def f(row: dict[str, str], key: str) -> float | None:
    value = row.get(key, "")
    if value == "" or value is None:
        return None
    return float(value)


def pct(value: float | None) -> str:
    if value is None:
        return "N/A"
    return f"{value * 100:.2f}%"


def num(value: float | None, digits: int = 4) -> str:
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}"


def ms(value: float | None) -> str:
    if value is None:
        return "N/A"
    return f"{value:.2f} ms"


def set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].font.name = "Microsoft YaHei"


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        document.add_paragraph(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr_cells[index].text = header
        for paragraph in hdr_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()


def rows_by_method_k(rows: list[dict[str, str]], *, group_type: str) -> dict[tuple[str, str], dict[str, str]]:
    return {
        (row["method"], row["k"]): row
        for row in rows
        if row["group_type"] == group_type
    }


def task_rows_by_method_task(rows: list[dict[str, str]]) -> dict[tuple[str, str], dict[str, str]]:
    return {
        (row["method"], row["task_type"]): row
        for row in rows
        if row["group_type"] == "task_type"
    }


def get_overall(overall_rows: list[dict[str, str]], method: str, k: str = "5") -> dict[str, str]:
    for row in overall_rows:
        if row["method"] == method and row["k"] == k:
            return row
    raise KeyError((method, k))


def get_task(task_rows: list[dict[str, str]], method: str, task_type: str, k: str = "5") -> dict[str, str]:
    for row in task_rows:
        if row["method"] == method and row["task_type"] == task_type and row["k"] == k:
            return row
    raise KeyError((method, task_type, k))


def build_failure_counts(failure_rows: list[dict[str, str]]) -> dict[str, Counter[str]]:
    counts: dict[str, Counter[str]] = defaultdict(Counter)
    for row in failure_rows:
        counts[row["method"]][row["task_type"]] += 1
        counts[row["method"]]["total"] += 1
    return counts


def build_oracle_selection_counts(results_dir: Path) -> Counter[str]:
    path = results_dir / "oracle_final_predictions.jsonl"
    if not path.exists():
        return Counter()
    return Counter(str(row.get("selected_oracle_method", "")) for row in read_jsonl(path))


def add_method_discussion(
    document: Document,
    method: str,
    overall_rows: list[dict[str, str]],
    task_rows: list[dict[str, str]],
) -> None:
    row = get_overall(overall_rows, method)
    semantic = get_task(task_rows, method, "semantic_fact")
    multihop = get_task(task_rows, method, "multi_hop_relation")
    exact = get_task(task_rows, method, "exact_file_lookup")
    document.add_heading(METHOD_LABELS[method], level=3)

    if method == "dense":
        add_paragraphs(
            document,
            [
                "Dense 基线使用 sentence-transformers/all-MiniLM-L6-v2 将问题和文档编码到同一个向量空间，然后用向量相似度排序文档。它代表语义检索能力：当问题经过改写、没有明显标题或精确词锚点时，向量模型可以通过语义接近性召回相关证据。",
                f"在整体 @5 上，Dense 的 Evidence Recall 为 {pct(f(row, 'evidence_recall'))}，Complete Evidence Recall 为 {pct(f(row, 'complete_evidence_recall'))}，MRR 为 {num(f(row, 'mrr'))}。三类任务中，semantic_fact 的 Complete@5 为 {pct(f(semantic, 'complete_evidence_recall'))}，exact_file_lookup 为 {pct(f(exact, 'complete_evidence_recall'))}，multi_hop_relation 为 {pct(f(multihop, 'complete_evidence_recall'))}。",
                "这个结果说明 Dense 很适合单事实语义问题，但在多跳问题上常常只能找到其中一个支持文档，难以保证 Top-k 同时覆盖完整证据链。对 exact/file lookup，它没有显式标题、短语和数字规则，因此不如专门的 File-FTS 稳定。",
            ],
        )
    elif method == "file_fts":
        add_paragraphs(
            document,
            [
                "File-FTS 基线使用 SQLite FTS5 建立全文检索索引。文档被切分为重叠 chunk，索引字段包括标题、正文、实体/关键词、数字日期等特征；查询时先在 chunk 层做 FTS/BM25 式召回，再聚合到文档层，并叠加标题匹配、精确短语、数字日期等重排序奖励。",
                f"在整体 @5 上，File-FTS 的 Evidence Recall 为 {pct(f(row, 'evidence_recall'))}，Complete Evidence Recall 为 {pct(f(row, 'complete_evidence_recall'))}，MRR 为 {num(f(row, 'mrr'))}。它在 exact_file_lookup 上的 Complete@5 达到 {pct(f(exact, 'complete_evidence_recall'))}，在 semantic_fact 上为 {pct(f(semantic, 'complete_evidence_recall'))}，但 multi_hop_relation 为 {pct(f(multihop, 'complete_evidence_recall'))}。",
                "这符合文件检索的预期优势：当问题包含标题、精确短语、日期或数字锚点时，File-FTS 可以快速定位文档；但多跳问题要求同时找齐多个支持文档，单纯词项匹配容易集中在第一个强相关文档上，完整证据覆盖不足。",
            ],
        )
    elif method == "graph_path":
        add_paragraphs(
            document,
            [
                "Graph-Path 基线是纯图检索方法。它把语料构造成文档-句子-实体-关系线索图：文档连接句子，句子连接实体，文档之间通过标题链接、共享实体、句级共现和关系线索建立可遍历路径。查询时抽取问题中的实体和关系词，作为起点进行 query-aware beam search，并根据路径覆盖、关系覆盖、多文档奖励和高频实体惩罚等因素给候选文档排序。",
                f"在整体 @5 上，Graph-Path 的 Evidence Recall 为 {pct(f(row, 'evidence_recall'))}，Complete Evidence Recall 为 {pct(f(row, 'complete_evidence_recall'))}，MRR 为 {num(f(row, 'mrr'))}。它在 multi_hop_relation 上的 Complete@5 为 {pct(f(multihop, 'complete_evidence_recall'))}，高于 Dense 和 File-FTS；在 exact_file_lookup 上为 {pct(f(exact, 'complete_evidence_recall'))}，semantic_fact 上为 {pct(f(semantic, 'complete_evidence_recall'))}。",
                "这说明图路径方法确实提供了互补能力：它不只是找最相似文本，而是试图沿实体和关系路径补齐第二跳证据。不过图扩展会引入噪声，所以 MRR 不一定最高，尤其在强标题/精确词任务上不如 File-FTS 直接。",
            ],
        )
    elif method == "fusion":
        add_paragraphs(
            document,
            [
                "Fusion 使用 Reciprocal Rank Fusion（RRF）融合 Dense、File-FTS 和 Graph-Path 三个检索结果。RRF 不依赖不同检索器的原始分数尺度，而是按每个文档在各列表中的排名累加 1/(k0 + rank)，因此适合融合向量、文件和图三类异构检索器。",
                f"Fusion 的整体 Complete@5 为 {pct(f(row, 'complete_evidence_recall'))}，高于三个单工具基线；Evidence@5 为 {pct(f(row, 'evidence_recall'))}，MRR 为 {num(f(row, 'mrr'))}。它的代价也最高：平均 Tool Calls 为 {num(f(row, 'average_tool_calls'), 2)}，平均延迟为 {ms(f(row, 'latency_avg_ms'))}。",
                "这说明三种检索器确实存在互补性：融合能把不同方法各自召回的证据放进同一个 Top-k，但它需要同时调用三个后端，成本明显高于单工具检索。后续智能体的目标不是简单复制 Fusion，而是在尽量少的调用次数下接近它的覆盖率。",
            ],
        )
    elif method == "oracle":
        add_paragraphs(
            document,
            [
                "Oracle 是 analysis-only 理想上限，不是可部署检索方法。它在评价阶段读取 gold labels，从 Dense、File-FTS、Graph-Path 和 Fusion 的已有输出中，为每个问题选择 Complete Recall 和 MRR 最优的那一个。它用于回答一个研究问题：如果路由器每题都能选对工具，最高能达到什么水平。",
                f"Oracle 的整体 Complete@5 为 {pct(f(row, 'complete_evidence_recall'))}，Evidence@5 为 {pct(f(row, 'evidence_recall'))}，MRR 为 {num(f(row, 'mrr'))}。这个结果高于 Fusion，说明仍存在进一步提升空间，主要来自更精确的问题级方法选择。",
                "由于 Oracle 使用评价标签选择结果，它不能参与实际部署，也不能用于调参。它只应作为上限参考，用来衡量后续 Rule Router 和 Two-Stage Agent 与理想选择之间的差距。",
            ],
        )


def create_report(
    *,
    baseline_dir: Path,
    results_dir: Path,
    figures_dir: Path,
    out_path: Path,
) -> None:
    overall_rows = read_csv(baseline_dir / "baseline_metrics_overall.csv")
    task_rows = read_csv(baseline_dir / "baseline_metrics_by_task.csv")
    all_rows = read_csv(baseline_dir / "baseline_metrics_all.csv")
    failure_rows = read_csv(baseline_dir / "failure_cases_k5.csv")
    failure_counts = build_failure_counts(failure_rows)
    oracle_counts = build_oracle_selection_counts(results_dir)

    document = Document()
    set_default_style(document)
    add_title(document, "冻结基线实验设计与结果分析报告")
    document.add_paragraph("数据集：final 冻结数据集；语料 3,500 篇文档；问题 180 个。")
    document.add_paragraph("方法：Dense / MiniLM、File-FTS、Graph-Path、Fusion / RRF、Oracle / 理想上限。")

    document.add_heading("1. 实验设定", level=1)
    add_paragraphs(
        document,
        [
            "本实验在 final 冻结数据集上进行。数据集包含 180 个问题，三类任务等量分布：semantic_fact、multi_hop_relation、exact_file_lookup 各 60 个。语料库固定为 3,500 篇文档，gold evidence 以外的文档均作为噪声或干扰文档。",
            "所有正式检索方法运行时只读取公开问题文本和公开语料内容，不读取 gold_documents、gold_sentences、答案、source id 或构造阶段元数据。预测文件写出以后，评价脚本才读取 gold labels 计算指标。这个流程保证最终测试集只用于评价，不用于方法调参。",
            "当前结果包括三种单检索基线、一个可运行融合基线，以及一个 analysis-only Oracle 上限。三种单检索基线已经冻结；Fusion 和 Oracle 用于补充分析方法互补性与理论空间。",
        ],
    )

    document.add_heading("2. 方法设计与原理", level=1)
    add_method_discussion(document, "dense", overall_rows, task_rows)
    add_method_discussion(document, "file_fts", overall_rows, task_rows)
    add_method_discussion(document, "graph_path", overall_rows, task_rows)
    add_method_discussion(document, "fusion", overall_rows, task_rows)
    add_method_discussion(document, "oracle", overall_rows, task_rows)

    document.add_heading("3. 指标定义", level=1)
    add_table(
        document,
        ["指标", "定义", "解释"],
        [
            ["Evidence Recall@k", "Top-k 是否至少包含一个 gold document", "衡量是否发现了基本证据"],
            ["Complete Evidence Recall@k", "Top-k 是否包含回答所需全部 gold documents", "多跳和综合检索的核心指标"],
            ["MRR", "第一个正确证据文档排名的倒数均值", "越高表示正确证据越靠前"],
            ["Search Success Rate", "本实验中等同于 Complete Evidence Recall@k", "衡量检索是否成功完成证据覆盖"],
            ["Average Tool Calls", "每题平均调用多少个核心检索后端", "直接反映检索成本"],
            ["Latency Avg / P95", "平均延迟与 95 分位延迟", "反映工程效率和稳定性"],
            ["Evidence Gain per Step", "第二轮相对第一轮新增证据", "当前单步基线不适用"],
            ["Stop Accuracy", "Judge 停止/继续判断准确率", "当前单步基线不适用"],
        ],
    )

    document.add_heading("4. Overall 结果", level=1)
    overall_by_method_k = rows_by_method_k(all_rows, group_type="overall")
    overall_table = []
    for method in METHOD_ORDER:
        for k in ("1", "3", "5"):
            row = overall_by_method_k[(method, k)]
            overall_table.append(
                [
                    METHOD_LABELS[method],
                    f"@{k}",
                    pct(f(row, "evidence_recall")),
                    pct(f(row, "complete_evidence_recall")),
                    num(f(row, "mrr")),
                    pct(f(row, "search_success_rate")),
                    num(f(row, "average_tool_calls"), 2),
                    ms(f(row, "latency_avg_ms")),
                    ms(f(row, "latency_p95_ms")),
                ]
            )
    add_table(
        document,
        ["方法", "k", "Evidence Recall", "Complete Recall", "MRR", "Search Success", "Avg Calls", "Avg Latency", "P95 Latency"],
        overall_table,
    )

    document.add_heading("5. 三类任务表现", level=1)
    complete_rows = []
    evidence_rows = []
    mrr_rows = []
    for method in METHOD_ORDER:
        complete_rows.append(
            [
                METHOD_LABELS[method],
                *[
                    pct(f(get_task(task_rows, method, task_type), "complete_evidence_recall"))
                    for task_type in TASK_ORDER
                ],
            ]
        )
        evidence_rows.append(
            [
                METHOD_LABELS[method],
                *[
                    pct(f(get_task(task_rows, method, task_type), "evidence_recall"))
                    for task_type in TASK_ORDER
                ],
            ]
        )
        mrr_rows.append(
            [
                METHOD_LABELS[method],
                *[
                    num(f(get_task(task_rows, method, task_type), "mrr"))
                    for task_type in TASK_ORDER
                ],
            ]
        )
    task_headers = ["方法", *[TASK_LABELS[task] for task in TASK_ORDER]]
    document.add_paragraph("Complete Evidence Recall@5：")
    add_table(document, task_headers, complete_rows)
    document.add_paragraph("Evidence Recall@5：")
    add_table(document, task_headers, evidence_rows)
    document.add_paragraph("MRR@5：")
    add_table(document, task_headers, mrr_rows)

    document.add_heading("6. 成本与效率", level=1)
    cost_rows = []
    for method in METHOD_ORDER:
        row = get_overall(overall_rows, method)
        cost_rows.append(
            [
                METHOD_LABELS[method],
                num(f(row, "average_tool_calls"), 2),
                ms(f(row, "latency_avg_ms")),
                ms(f(row, "latency_p95_ms")),
                "analysis-only，不可部署" if method == "oracle" else "可运行方法",
            ]
        )
    add_table(document, ["方法", "Avg Tool Calls", "Avg Latency", "P95 Latency", "成本口径"], cost_rows)
    add_paragraphs(
        document,
        [
            "File-FTS 的平均延迟最低，适合精确定位和大规模快速检索。Graph-Path 延迟高于 File-FTS，但仍明显低于远程 Dense。Dense 的延迟主要来自远程 embedding 调用和向量检索流程。Fusion 需要同时调用三个检索器，因此完整召回率提升的同时也付出了最高调用成本和最高端到端延迟。",
            "Oracle 的成本不是部署成本。它选择的是已有输出中最好的一个，用于分析理想路由上限；因此不能用 Oracle 的延迟和调用数评价真实系统效率。",
        ],
    )

    document.add_heading("7. 失败案例统计", level=1)
    failure_table = []
    for method in METHOD_ORDER:
        counts = failure_counts[method]
        failure_table.append(
            [
                METHOD_LABELS[method],
                str(counts["total"]),
                str(counts["semantic_fact"]),
                str(counts["multi_hop_relation"]),
                str(counts["exact_file_lookup"]),
            ]
        )
    add_table(document, ["方法", "Complete@5 失败总数", "Semantic Fact", "Multi-hop Relation", "Exact/File Lookup"], failure_table)
    add_paragraphs(
        document,
        [
            "失败统计显示，多跳关系问题仍是最主要瓶颈。File-FTS 和 Dense 在 multi_hop_relation 中经常能召回至少一个相关文档，但 Complete Recall 低，说明它们缺少稳定补齐第二跳证据的机制。Graph-Path 在多跳任务上更强，但仍会受到实体抽取噪声、关系路径歧义和高频实体扩散影响。",
            "Exact/File Lookup 的失败主要来自标题不完全匹配、短语在噪声文档中重复出现、或者数字日期锚点不足以唯一定位。Semantic Fact 的失败更多来自语义改写后问题与原始文档表达之间的距离，以及 Dense/File/Graph 对证据排序的差异。",
        ],
    )

    if oracle_counts:
        document.add_heading("8. Oracle 选择分布", level=1)
        oracle_table = [
            [METHOD_LABELS.get(method, method), str(count)]
            for method, count in sorted(oracle_counts.items(), key=lambda item: (-item[1], item[0]))
        ]
        add_table(document, ["Oracle 选择的底层输出", "问题数"], oracle_table)
        add_paragraphs(
            document,
            [
                "Oracle 选择分布可以反映三种检索方式和 Fusion 的互补性。如果 Oracle 经常选择不同方法，说明不存在单一检索器可以稳定覆盖全部任务；后续 Router 或 Two-Stage Agent 的价值就在于根据问题特征选择更合适的检索路径。",
            ],
        )

    document.add_heading("9. 总结性评述", level=1)
    graph = get_overall(overall_rows, "graph_path")
    fusion = get_overall(overall_rows, "fusion")
    oracle = get_overall(overall_rows, "oracle")
    add_paragraphs(
        document,
        [
            f"从单工具基线看，File-FTS 的 MRR 和精确检索表现最好，Graph-Path 的 multi-hop Complete@5 最强，Dense 在 semantic_fact 上稳定。这说明三类检索方式确实对应了不同知识检索形态：语义相似、文件/精确定位、实体关系路径。",
            f"从整体 Complete@5 看，Fusion 达到 {pct(f(fusion, 'complete_evidence_recall'))}，高于最强单工具 Graph-Path 的 {pct(f(graph, 'complete_evidence_recall'))}，说明多索引融合有明确收益。但 Fusion 平均调用 3 个后端，成本高，不适合作为轻量智能体的最终目标。",
            f"Oracle 的 Complete@5 为 {pct(f(oracle, 'complete_evidence_recall'))}，高于 Fusion，说明如果能更准确地判断问题类型和证据缺口，还有进一步提升空间。后续研究应围绕两个方向展开：一是 Rule Router 是否能用低成本规则选择合适工具；二是 Two-Stage Agent 是否能在第一轮证据不足时追加最有价值的第二个工具，从而以低于 Fusion 的成本接近 Oracle 上限。",
            "因此，本阶段基线已经形成清晰、可比较、可解释的实验底座。后续不应再根据 final 结果调整这五个结果的参数；新的 Router、Judge 或 Agent 设计需要在隔离的验证材料上确定，再回到 final 数据集做一次性评价。",
        ],
    )

    dashboard = figures_dir / "summary" / "baseline_evaluation_dashboard.png"
    if dashboard.exists():
        document.add_heading("附图：基线可视化总览", level=1)
        document.add_picture(str(dashboard), width=Inches(6.2))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Create a DOCX baseline report.")
    parser.add_argument("--baseline-dir", default=str(DEFAULT_BASELINE_DIR))
    parser.add_argument("--results-dir", default=str(DEFAULT_RESULTS_DIR))
    parser.add_argument("--figures-dir", default=str(DEFAULT_FIGURES_DIR))
    parser.add_argument("--out", default=str(DEFAULT_REPORT_PATH))
    return parser


def main() -> int:
    args = build_arg_parser().parse_args()
    create_report(
        baseline_dir=Path(args.baseline_dir),
        results_dir=Path(args.results_dir),
        figures_dir=Path(args.figures_dir),
        out_path=Path(args.out),
    )
    print(args.out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
