from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = PROJECT_ROOT / "docs" / "reports" / "最终项目报告_轻量级多索引智能体搜索.docx"

DATA_DIR = PROJECT_ROOT / "data" / "final"
RESULTS_DIR = PROJECT_ROOT / "results" / "final"
VIS_DIR = PROJECT_ROOT / "figures" / "agent"
MANIFEST_PATH = (
    PROJECT_ROOT
    / "results"
    / "analysis"
    / "agent"
    / "freeze_manifest.original-layout.json"
)

METHOD_ORDER = [
    "dense",
    "file_fts",
    "graph_path",
    "rule_router",
    "fusion",
    "oracle",
    "two_stage_agent",
]
METHOD_NAME = {
    "dense": "Dense",
    "file_fts": "File FTS",
    "graph_path": "Graph Path",
    "rule_router": "Rule Router",
    "fusion": "Fusion",
    "oracle": "Oracle",
    "two_stage_agent": "Two-Stage Agent",
}
TASK_ORDER = ["semantic_fact", "multi_hop_relation", "exact_file_lookup"]
TASK_NAME = {
    "semantic_fact": "语义事实问题",
    "multi_hop_relation": "多跳关系问题",
    "exact_file_lookup": "精确文件定位问题",
}


def main() -> None:
    rows = load_results()
    manifest = load_manifest()
    dataset_stats = load_dataset_stats()
    doc = build_report(rows, manifest, dataset_stats)
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(f"generated: {REPORT_PATH}")


def load_results() -> list[dict[str, str]]:
    with (RESULTS_DIR / "main_results.csv").open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def load_manifest() -> dict:
    if MANIFEST_PATH.exists():
        return json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    return {}


def load_dataset_stats() -> dict[str, object]:
    questions_path = DATA_DIR / "questions.jsonl"
    corpus_path = DATA_DIR / "corpus.jsonl"
    task_counts: dict[str, int] = {}
    gold_counts: list[int] = []
    with questions_path.open("r", encoding="utf-8") as f:
        for line in f:
            q = json.loads(line)
            task_counts[q["task_type"]] = task_counts.get(q["task_type"], 0) + 1
            gold_counts.append(len(q.get("gold_documents", [])))
    corpus_count = sum(1 for _ in corpus_path.open("r", encoding="utf-8"))
    return {
        "question_count": sum(task_counts.values()),
        "corpus_count": corpus_count,
        "task_counts": task_counts,
        "avg_gold_docs": sum(gold_counts) / len(gold_counts),
        "max_gold_docs": max(gold_counts),
    }


def build_report(rows: list[dict[str, str]], manifest: dict, dataset_stats: dict[str, object]) -> Document:
    doc = Document()
    configure_document(doc)

    add_title(doc, "轻量级多索引智能体搜索项目最终报告")
    add_center(doc, "Lightweight Task-Aware Multi-Index Agentic Search")
    add_center(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_para(
        doc,
        "本报告汇总整个项目的最终数据集、三类核心检索基线、路由器、两阶段智能体搜索方法、实验结果和可视化分析。报告中的实验结果均来自 results/final，最终智能体方法统一命名为 two_stage_agent。",
    )

    add_heading(doc, "摘要", 1)
    agent_c5 = metric(rows, "two_stage_agent", 5, "overall", "overall", "complete_recall_at_k")
    fusion_c5 = metric(rows, "fusion", 5, "overall", "overall", "complete_recall_at_k")
    agent_calls = metric(rows, "two_stage_agent", 5, "overall", "overall", "tool_calls")
    fusion_calls = metric(rows, "fusion", 5, "overall", "overall", "tool_calls")
    router_c5 = metric(rows, "rule_router", 5, "overall", "overall", "complete_recall_at_k")
    add_para(
        doc,
        f"项目最终构建了一个轻量级多索引智能体检索系统。系统不固定调用所有检索器，而是先由 Rule Router 根据问题类型选择最合适的首轮检索器，再由 Evidence Judge 判断 Top-k 证据是否足够，必要时调用第二个检索器，并通过 protected weighted RRF 融合结果。最终 two_stage_agent 的 Complete@5 为 {pct_value(agent_c5)}，与全量 Fusion 的 {pct_value(fusion_c5)} 持平；但 Agent 平均只调用 {agent_calls:.2f} 个工具，而 Fusion 固定调用 {fusion_calls:.2f} 个工具。相比 Rule Router 的 Complete@5={pct_value(router_c5)}，Agent 进一步提升了完整证据召回，同时保持较低检索成本。",
    )

    add_heading(doc, "1. 研究目标与问题定义", 1)
    add_para(
        doc,
        "本项目研究的问题是：在同一个大规模文档库中，不同问题需要不同检索能力。语义事实问题更适合向量语义检索，精确文件定位问题更适合文件/标题/短语检索，多跳关系问题更需要实体关系和路径结构。若固定使用单一检索器，会在某些任务上表现不足；若固定融合所有检索器，成本较高。因此，本项目尝试构建一个轻量级、任务感知、可解释的多索引智能体搜索框架。",
    )
    add_para(
        doc,
        "核心研究假设是：通过问题路由、证据充分性判断和按需二轮检索，可以在不固定调用所有工具的情况下，接近全量融合的完整证据召回能力，并在三类任务之间取得更好的平衡。",
    )

    add_heading(doc, "2. 冻结数据集设计", 1)
    add_para(
        doc,
        "最终数据集已经冻结，后续实验不再修改。所有方法均从同一个 3500 篇文档的公开语料库中检索；gold documents 只在评价阶段使用，不进入检索器、路由器或智能体运行过程。",
    )
    add_table(
        doc,
        ["项目", "数值"],
        [
            ["文档总数", f"{dataset_stats['corpus_count']} 篇"],
            ["问题总数", f"{dataset_stats['question_count']} 题"],
            ["semantic_fact", f"{dataset_stats['task_counts'].get('semantic_fact', 0)} 题"],
            ["multi_hop_relation", f"{dataset_stats['task_counts'].get('multi_hop_relation', 0)} 题"],
            ["exact_file_lookup", f"{dataset_stats['task_counts'].get('exact_file_lookup', 0)} 题"],
            ["平均 gold documents 数", f"{float(dataset_stats['avg_gold_docs']):.2f}"],
            ["最大 gold documents 数", str(dataset_stats["max_gold_docs"])],
        ],
    )
    add_para(
        doc,
        "三类任务分别承担不同实验目的：semantic_fact 主要检验语义匹配能力；multi_hop_relation 检验完整多跳证据发现能力；exact_file_lookup 保留文件检索优势，但不完全依赖标题，使文件索引有明确发挥空间。数据集冻结后，Router、Judge 和参数不再使用 final 数据继续调参，以避免测试集污染。",
    )
    add_figure(
        doc,
        PROJECT_ROOT / "docs" / "assets" / "最终数据集结构图.png",
        "图 1  最终数据集结构示意图。",
        width=6.6,
    )

    add_heading(doc, "3. 方法体系", 1)
    add_heading(doc, "3.1 三类核心检索器", 2)
    add_table(
        doc,
        ["方法", "核心实现", "优势", "不足"],
        [
            ["Dense", "sentence-transformers/all-MiniLM-L6-v2 向量检索。", "适合语义事实、同义改写和非精确表达。", "多跳问题中常找到部分证据，但不保证完整证据组合。"],
            ["File FTS", "SQLite FTS5 + BM25F 风格打分，结合 chunk/doc aggregation、标题权重、短语和数字日期锚点。", "适合标题、文件、短语、编号、日期等精确定位。", "多跳任务中 Evidence 高但 Complete 低，缺少关系补全能力。"],
            ["Graph Path", "PathRetriever-style 图路径检索，基于实体-文档-句子/标题图和 beam search。", "适合实体关系、多跳桥接和证据组合。", "对事实题和文件定位题不如 file/dense 稳定，依赖实体抽取质量。"],
        ],
    )

    add_heading(doc, "3.2 Rule Router", 2)
    add_para(
        doc,
        "Rule Router 是第一阶段工具选择器。它从问题文本中抽取 file cues、lookup cues、weak exact cues、graph cues、bridge patterns、quoted phrases、numbers/dates 和 simple entities，并据此计算 file_score、graph_score、dense_score。强文件定位信号优先 file_fts，强实体关系或桥接信号优先 graph_path，其余默认 dense。",
    )
    add_para(
        doc,
        "Rule Router 的价值在于用一次可解释的轻量选择减少无差别多工具调用。但它只能选择一个首轮工具，无法判断检索结果是否已经覆盖完整证据，因此还需要 Evidence Judge。",
    )

    add_heading(doc, "3.3 Evidence Judge 与 Two-Stage Agent", 2)
    add_para(
        doc,
        "Two-Stage Agent 的流程为：第一步由 Rule Router 选择首轮工具并返回 Top-k；第二步由 Evidence Judge 判断当前 Top-k 是否足以覆盖问题所需证据；若不足，则选择第二个工具补检索；最后使用 protected weighted RRF 融合两轮结果。",
    )
    add_para(
        doc,
        "这里的“判断 Top-k 是否足以覆盖问题所需证据”不是使用 gold documents，也不是生成答案，而是使用运行时可见的结构化信号进行启发式判断。Judge 观察 top score、score gap、title keyword overlap、keyword coverage ratio、entity hits、entity coverage ratio、exact anchor in top3、dense_confident、semantic_confident、graph_need、file_need 和 graph_sufficient 等信号。如果问题呈现多实体或桥接结构，但当前 Top-k 只覆盖部分实体或分数不够稳定，则认为可能缺多跳证据，转向 graph_path；如果问题有标题、引号、数字日期等文件锚点，但当前结果未稳定命中，则转向 file_fts；如果 dense 结果置信度高且覆盖了主要关键词/实体，则停止。",
    )
    add_table(
        doc,
        ["首轮工具", "停止依据", "继续依据", "二轮工具"],
        [
            ["dense", "分数高、gap 明显、语义覆盖足够，且无明显 graph/file 风险。", "实体覆盖不足、多跳桥接风险、或文件锚点未验证。", "graph_path 或 file_fts"],
            ["file_fts", "文件锚点、标题、短语或数字日期在 Top-3 中已稳定出现。", "文件结果只覆盖局部证据，仍存在关系/桥接风险。", "graph_path 或 dense"],
            ["graph_path", "图信号强，实体覆盖或关键词覆盖足够。", "图检索覆盖不足，或问题实为文件定位。", "dense 或 file_fts"],
        ],
    )
    add_para(
        doc,
        "最终融合采用 protected weighted RRF。普通 RRF 可能把第一轮高置信证据挤出 Top-k；protected weighted RRF 会保护第一轮高置信 Top-1，同时在 graph_need 或 file_need 明确时提高二轮工具权重，从而兼顾排序稳定性和补证据能力。",
    )

    add_heading(doc, "4. 评价指标", 1)
    add_table(
        doc,
        ["指标", "定义", "意义"],
        [
            ["Evidence Recall@k", "Top-k 至少包含一个 gold document 的问题比例。", "衡量基本证据发现。"],
            ["Complete Evidence Recall@k", "Top-k 包含回答所需全部 gold documents 的问题比例。", "衡量完整证据发现，是多跳任务核心指标。"],
            ["MRR", "第一个正确证据排名的倒数均值。", "衡量排序质量。"],
            ["Search Success@k", "本实验中等同于 Complete Evidence Recall@k。", "表示完整检索成功率。"],
            ["Average Tool Calls", "每题平均调用核心检索器数量。", "衡量检索成本。"],
            ["Latency / P95 Latency", "平均与 P95 时延。", "反映工程效率，受远程 embedding 调用影响较大。"],
            ["Evidence Gain per Step", "第二轮新增 gold evidence 的比例。", "衡量二轮检索收益。"],
            ["Stop Accuracy", "证据充分时停、缺证据时继续的比例。", "分析 Judge 决策质量。"],
        ],
    )

    add_heading(doc, "5. 总体实验结果", 1)
    add_para(doc, "下图展示 @5 下各方法在 Evidence Recall、Complete Evidence Recall 和 MRR 上的整体表现。")
    add_figure(
        doc,
        VIS_DIR / "00_overview" / "overall_metrics_at5.png",
        "图 2  总体检索性能 @5。",
        width=7.0,
    )
    add_para(doc, "下表给出 @3 和 @5 的总体指标。Oracle 是使用 gold 信息的分析上界，不是可部署方法。")
    add_overall_table(doc, rows, 3)
    add_overall_table(doc, rows, 5)
    add_para(
        doc,
        "从总体结果看，two_stage_agent 的 Complete@5 为 85.56%，与固定三工具 Fusion 持平；但 Agent 的平均工具调用为 1.44，明显低于 Fusion 的 3.00。相较三个单一基线，Agent 在完整证据召回上更强；相较 Rule Router，Agent 通过二轮证据补检索将 Complete@5 从 82.22% 提升到 85.56%。",
    )
    add_figure(
        doc,
        VIS_DIR / "00_overview" / "complete_recall_at3_at5.png",
        "图 3  Complete Evidence Recall 在 @3 与 @5 下的对比。",
        width=6.8,
    )

    add_heading(doc, "6. 三类任务上的表现差异", 1)
    add_para(doc, "三类任务体现了不同检索形式的优缺点。下图展示 Complete@5 在三类任务上的热力图。")
    add_figure(
        doc,
        VIS_DIR / "01_task_breakdown" / "complete_recall_heatmap_at5.png",
        "图 4  三类任务 Complete Evidence Recall@5 热力图。",
        width=6.4,
    )
    add_task_ec_table(doc, rows, 3)
    add_task_ec_table(doc, rows, 5)
    add_para(
        doc,
        "File FTS 在 semantic_fact 和 exact_file_lookup 上表现很强，但在 multi_hop_relation 上 Complete@5 只有 41.67%，说明文档/文件检索可以找到相关证据，却难以补齐多跳证据组合。Graph Path 在 multi_hop_relation 上 Complete@5 达到 63.33%，但在 semantic_fact 和 exact_file_lookup 上不如 File FTS 稳定。two_stage_agent 通过路由和证据判断，保留了 file/dense 在简单任务上的稳定性，并在多跳任务上利用 graph_path 补证据，实现了更均衡的任务表现。",
    )
    add_figure(
        doc,
        VIS_DIR / "03_agent_analysis" / "core_retriever_complementarity_at5.png",
        "图 5  三类核心检索器与 Agent 的互补性。",
        width=6.8,
    )

    add_heading(doc, "7. 成本与效率分析", 1)
    add_para(doc, "Agent 的关键优势不是在所有指标上超过所有方法，而是在较低工具调用成本下接近或达到全量融合的完整证据召回。")
    add_figure(
        doc,
        VIS_DIR / "02_cost_efficiency" / "cost_effectiveness_complete_at5.png",
        "图 6  Complete@5 与平均工具调用次数的成本-效果关系。",
        width=6.4,
    )
    add_figure(
        doc,
        VIS_DIR / "02_cost_efficiency" / "taskwise_tool_calls_at5.png",
        "图 7  三类任务上的平均工具调用次数。",
        width=6.8,
    )
    add_para(
        doc,
        "从成本看，Fusion 固定调用三个检索器，因此 Complete@5 很强但调用成本高。two_stage_agent 只在 Judge 判断证据不足时进行第二轮调用，平均调用 1.44 次，能在保持 Complete@5=85.56% 的同时显著降低工具调用数。Latency 当前受 DeepInfra 远程 embedding 调用影响较大，因此报告更推荐使用 Average Tool Calls 作为方法级成本指标。",
    )

    add_heading(doc, "8. Agent 专项分析", 1)
    add_figure(
        doc,
        VIS_DIR / "03_agent_analysis" / "agent_task_profile_at5.png",
        "图 8  最终 Agent 在三类任务上的 Evidence、Complete 和 MRR 画像。",
        width=6.6,
    )
    add_figure(
        doc,
        VIS_DIR / "03_agent_analysis" / "task_balance_complete_at5.png",
        "图 9  任务均衡性：平均 Complete@5 与最差任务 Complete@5。",
        width=6.2,
    )
    add_para(
        doc,
        "Agent 的价值体现在平衡三类检索器：Dense 解决语义表达，File FTS 保留文件定位优势，Graph Path 弥补多跳关系证据。Rule Router 负责第一轮选择，Evidence Judge 决定是否补检索，protected RRF 防止二轮结果破坏首轮高置信证据。这使系统能在不固定全量融合的情况下获得接近 Fusion 的完整证据召回。",
    )

    add_heading(doc, "9. 失败模式与系统边界", 1)
    add_table(
        doc,
        ["问题", "表现", "原因", "后续方向"],
        [
            ["多跳 Top-3 仍不够强", "multi_hop_relation Complete@3 为 46.67%。", "二轮能补到证据，但完整证据组合不总能进入 Top-3。", "在验证集上研究证据组合重排或路径级 rerank。"],
            ["Exact File Lookup 有过度二轮调用", "exact_file_lookup Avg Calls 较高。", "Judge 对部分文件题过于谨慎。", "增强文件锚点停止规则，但不能在 final 上调参。"],
            ["图检索依赖实体抽取", "Graph Path 对事实题和文件题不稳定。", "实体别名、简称和跨句关系抽取可能不完整。", "改进实体规范化和句级共现图。"],
            ["远程 dense 带来时延", "包含 dense 的方法 P95 latency 较高。", "DeepInfra embedding 请求是主要工程开销。", "使用本地模型、缓存 query embedding、批处理。"],
            ["Oracle 仍高于 Agent", "Oracle Complete@5 为 90.00%。", "Oracle 使用 gold 信息，是不可部署上界。", "说明仍有改进空间，但不能用 final 继续调参。"],
        ],
    )

    add_heading(doc, "10. 最终结论", 1)
    add_para(
        doc,
        "本项目完成了从数据集构建、三类基线实现、路由器设计、Evidence Judge、两阶段 Agent 到冻结实验和可视化报告的完整闭环。最终系统证明：多索引检索不必固定全量融合，通过任务感知路由和证据充分性判断，可以在较低成本下获得接近全量融合的完整证据召回。",
    )
    add_para(
        doc,
        "最终 two_stage_agent 在 Complete@5 上达到 85.56%，与 Fusion 持平，但平均工具调用只有 1.44 次；同时它在 semantic_fact、multi_hop_relation、exact_file_lookup 三类任务上比单一检索器更均衡。该结果支持项目的核心结论：轻量级任务感知多索引智能体搜索能够在控制成本的同时平衡语义检索、文件检索和图检索的优缺点。",
    )
    add_para(
        doc,
        "后续工作应遵守实验隔离原则，不再用 final 数据调参。可以在独立验证集上继续研究多跳证据组合重排、文件题停止规则、图实体抽取增强和 dense 缓存/本地化部署。",
    )

    add_heading(doc, "附录：文件位置", 1)
    add_table(
        doc,
        ["内容", "路径"],
        [
            ["冻结数据集", "data/final"],
            ["冻结实验结果", "results/final"],
            ["可视化图表", "figures/agent"],
            ["最终 Agent 代码", "src/agent.py"],
            ["Router 代码", "src/router.py"],
            ["Evidence Judge 代码", "src/evidence_judge.py"],
            ["历史冻结清单", "results/analysis/agent/freeze_manifest.original-layout.json"],
        ],
    )

    return doc


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 19, True)


def add_center(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 15 if level == 1 else 12.5, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, 8.8, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, 8.2)
    doc.add_paragraph()


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.5) -> None:
    if not image_path.exists():
        add_para(doc, f"图像缺失：{image_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, 9.2, True)


def add_overall_table(doc: Document, rows: list[dict[str, str]], k: int) -> None:
    table_rows = []
    for method in METHOD_ORDER:
        row = find_row(rows, method, k, "overall", "overall")
        table_rows.append(
            [
                METHOD_NAME[method],
                pct(row["evidence_recall_at_k"]),
                pct(row["complete_recall_at_k"]),
                pct(row["mrr"]),
                f'{float(row["tool_calls"]):.2f}',
            ]
        )
    add_table(doc, ["方法", f"Evidence@{k}", f"Complete@{k}", "MRR", "Avg Calls"], table_rows)


def add_task_ec_table(doc: Document, rows: list[dict[str, str]], k: int) -> None:
    table_rows = []
    for method in METHOD_ORDER:
        table_rows.append(
            [
                METHOD_NAME[method],
                ec_cell(rows, method, k, "semantic_fact"),
                ec_cell(rows, method, k, "multi_hop_relation"),
                ec_cell(rows, method, k, "exact_file_lookup"),
            ]
        )
    add_table(doc, ["方法", f"semantic_fact E/C@{k}", f"multi_hop E/C@{k}", f"exact_file E/C@{k}"], table_rows)


def find_row(rows: list[dict[str, str]], method: str, k: int, group_type: str, group_name: str) -> dict[str, str]:
    for row in rows:
        if row["method"] == method and int(row["k"]) == k and row["group_type"] == group_type and row["group_name"] == group_name:
            return row
    raise KeyError((method, k, group_type, group_name))


def metric(rows: list[dict[str, str]], method: str, k: int, group_type: str, group_name: str, key: str) -> float:
    return float(find_row(rows, method, k, group_type, group_name)[key])


def pct(value: str) -> str:
    return pct_value(float(value))


def pct_value(value: float) -> str:
    return f"{value * 100:.2f}%"


def ec_cell(rows: list[dict[str, str]], method: str, k: int, task: str) -> str:
    row = find_row(rows, method, k, "task_type", task)
    return f'{pct(row["evidence_recall_at_k"])} / {pct(row["complete_recall_at_k"])}'


if __name__ == "__main__":
    main()
