from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results" / "final"
REPORT_ROOT = ROOT / "router_and_agent_design_and_results.docx"
REPORT_DIR = ROOT / "reports" / "router_and_agent_design_and_results.docx"

METHOD_ORDER = [
    "dense",
    "file_fts",
    "graph_path",
    "rule_router",
    "fusion",
    "oracle",
    "two_stage_agent",
]
METHOD_NAME = {
    "dense": "Dense",
    "file_fts": "File FTS",
    "graph_path": "Graph Path",
    "rule_router": "Rule Router",
    "fusion": "Fusion",
    "oracle": "Oracle",
    "two_stage_agent": "Two-Stage Agent",
}
TASK_ORDER = ["semantic_fact", "multi_hop_relation", "exact_file_lookup"]


def main() -> None:
    rows = load_main_results()
    manifest = load_manifest()
    doc = build_report(rows, manifest)
    REPORT_ROOT.parent.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_ROOT)
    doc.save(REPORT_DIR)
    print(f"已生成：{REPORT_ROOT}")
    print(f"已生成：{REPORT_DIR}")


def load_main_results() -> list[dict[str, str]]:
    with (RESULTS_DIR / "main_results.csv").open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def load_manifest() -> dict:
    path = ROOT / "manifest" / "freeze_manifest.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def build_report(rows: list[dict[str, str]], manifest: dict) -> Document:
    doc = Document()
    configure_document(doc)

    add_title(doc, "路由与两阶段智能体搜索：设计、实现与冻结实验结果")
    add_center(doc, "Frozen package: router_and_agent")
    add_center(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_para(
        doc,
        "本文档记录最终冻结的 Router + Two-Stage Agent 检索系统。冻结包只保留三个核心基线 Dense、File FTS、Graph Path，以及 Rule Router、Fusion、Oracle 和最终 two_stage_agent。最终智能体统一命名为 two_stage_agent。",
    )

    add_heading(doc, "1. 冻结对象与实验隔离", 1)
    add_para(
        doc,
        "本次冻结的目标是将数据集、实现代码、实验输出和报告放入一个独立归档目录，作为后续论文、汇报和可视化的唯一依据。所有方法在同一冻结 final 数据集上运行，运行时只读取公开问题和公开文档；gold documents 只在评价阶段由 ExperimentBox 读取。",
    )
    dataset = manifest.get("frozen_dataset", {})
    task_counts = dataset.get("task_counts", {})
    add_table(
        doc,
        ["项目", "冻结值"],
        [
            ["语料规模", f"{dataset.get('corpus_document_count', 3500)} 篇文档"],
            ["问题规模", f"{dataset.get('question_count', 180)} 个问题"],
            ["semantic_fact", f"{task_counts.get('semantic_fact', 60)} 题"],
            ["multi_hop_relation", f"{task_counts.get('multi_hop_relation', 60)} 题"],
            ["exact_file_lookup", f"{task_counts.get('exact_file_lookup', 60)} 题"],
            ["正式智能体名称", "two_stage_agent"],
            ["保留方法", "dense, file_fts, graph_path, rule_router, fusion, oracle, two_stage_agent"],
        ],
    )
    add_para(
        doc,
        "隔离原则：检索方法只能访问 data/final/corpus.jsonl 与公开问题字段；评价指标从 data/final/questions.jsonl 中读取 gold labels。DeepInfra 密钥只通过 DEEPINFRA_TOKEN 环境变量使用，归档中不保存密钥值。",
    )

    add_heading(doc, "2. 三类核心搜索器", 1)
    add_para(
        doc,
        "系统保留三类互补搜索器：Dense 负责语义相似性，File FTS 负责文件级、标题级和精确锚点检索，Graph Path 负责实体关系和多跳路径检索。三者分别代表语义向量检索、文档/文件倒排检索、结构化图检索三种常见知识检索形式。",
    )
    add_table(
        doc,
        ["方法", "实现方式", "优势", "典型不足"],
        [
            ["Dense", "MiniLM 句向量检索；对问题和文档编码后做向量相似度排序。", "适合语义事实、同义改写、非精确表达。", "多跳问题中容易找到部分证据，但难以保证完整证据组合。"],
            ["File FTS", "SQLite FTS5 + BM25F 风格打分；结合 chunk/doc aggregation、标题权重、exact phrase、数字日期锚点。", "适合文件定位、标题定位、精确短语和编号日期检索。", "在 multi-hop 中 Evidence 可能很高，但 Complete 较低，因为缺少关系补全能力。"],
            ["Graph Path", "PathRetriever-style 图路径检索；构建实体-文档-句子/标题链接，用 beam search 查找关系路径。", "适合实体关联、多跳桥接和证据组合。", "在简单事实和精确文件题上可能不如 file/dense 稳，且实体抽取质量影响较大。"],
        ],
    )

    add_heading(doc, "3. Rule Router 设计", 1)
    add_para(
        doc,
        "Rule Router 是第一阶段工具选择器。它不读取 gold labels，只从问题文本提取可解释信号，然后在 dense、file_fts、graph_path 中选择一个首轮工具。",
    )
    add_para(
        doc,
        "主要信号包括：file cues，如 file、document、article、record；lookup cues，如 find、locate、title、mentions；weak exact cues，如 year、date、number；graph cues，如 born、directed、written、located、whose；bridge patterns，如 whose、which ... was/is、person who、in which；以及 simple entities、quoted phrases、numbers/dates。",
    )
    add_para(
        doc,
        "路由分数由 file_score、graph_score、dense_score 三部分组成。强文件定位问题优先 file_fts；强实体关系或桥接问题优先 graph_path；没有明显文件或图信号的语义事实问题默认 dense。Rule Router 的作用不是最终答案判断，而是用一次廉价、可解释的工具选择降低无差别多工具调用成本。",
    )

    add_heading(doc, "4. Two-Stage Agent 设计", 1)
    add_para(
        doc,
        "最终 two_stage_agent 在 Rule Router 之后增加 Evidence Judge。整体流程为：第一轮路由并调用一个核心搜索器；Evidence Judge 判断当前证据是否足够；若不足，最多调用第二个搜索器；最后用 protected weighted RRF 生成最终 Top-k。",
    )
    add_para(
        doc,
        "Evidence Judge 的核心不是预测答案，而是判断当前 Top-k 是否足以覆盖问题所需证据。它使用以下结构化信号：top score、score gap、title keyword overlap、keyword coverage ratio、entity hits、entity coverage ratio、exact anchor in top3、dense_confident、semantic_confident、graph_need、file_need、graph_sufficient。",
    )
    add_table(
        doc,
        ["首轮工具", "停止条件", "继续条件", "二轮工具倾向"],
        [
            ["dense", "dense 置信度足够，且没有明显 file_need 或 graph_need。", "存在文件定位锚点，或实体/桥接覆盖不足。", "file_need -> file_fts；graph_need -> graph_path。"],
            ["file_fts", "file_score 高，Top-3 已覆盖 exact/title/number anchor，且无多跳风险。", "文件结果只覆盖局部证据，问题仍呈现关系/桥接需求。", "优先 graph_path；无法锚定时用 dense 交叉验证。"],
            ["graph_path", "图信号强，实体覆盖或关键词覆盖足够。", "图检索覆盖不足，或实际包含强文件定位信号。", "file_need -> file_fts；否则 dense 补语义证据。"],
        ],
    )
    add_para(doc, "二轮查询对 graph_path 会附加首轮 Top-2 标题，帮助图检索获得更明确的起点；对 dense/file_fts 保持问题文本稳定，避免引入过多噪声。")

    add_heading(doc, "5. Protected Weighted RRF", 1)
    add_para(
        doc,
        "普通 RRF 能融合两轮结果，但可能把第一轮高置信正确证据挤出 Top-k。最终系统使用 protected weighted RRF：先根据 Judge 信号保护第一轮 Top-1，再对二轮工具按需求加权。例如 graph_need 时提高 graph_path 权重，file_need 时提高 file_fts 权重，dense_confident 或 exact_anchor 时保护第一轮结果。",
    )
    add_para(
        doc,
        "该策略的目标是同时获得两类收益：保留首轮高质量证据的排序稳定性；让二轮结果补充缺失实体、缺失关系或文件锚点，从而提升 Complete Evidence Recall。",
    )

    add_heading(doc, "6. 指标定义", 1)
    add_table(
        doc,
        ["指标", "定义", "解释"],
        [
            ["Evidence Recall@k", "Top-k 中至少包含一个 gold document 的问题比例。", "衡量是否发现了基本证据。"],
            ["Complete Evidence Recall@k", "Top-k 中包含回答所需全部 gold documents 的问题比例。", "多跳和综合检索的核心指标。"],
            ["Search Success@k", "本实验中等同于 Complete Evidence Recall@k。", "表示完整证据检索成功率。"],
            ["MRR", "第一个正确证据排名的倒数均值。", "衡量排序质量。"],
            ["Average Tool Calls", "平均每题调用核心检索工具数量。", "衡量搜索成本。"],
            ["Latency / P95 Latency", "端到端平均与 P95 时延。", "工程效率指标；远程 dense 调用会显著影响该值。"],
            ["Evidence Gain per Step", "第二轮新增 gold evidence 的比例。", "衡量二轮检索是否有实际收益。"],
            ["Stop Accuracy", "证据充分时停、缺证据时继续的比例。", "分析 Evidence Judge 的决策质量。"],
        ],
    )

    add_heading(doc, "7. 总体实验结果", 1)
    for k in [3, 5]:
        add_para(doc, f"总体 @{k}：")
        table_rows = []
        for method in METHOD_ORDER:
            r = find_row(rows, method, k, "overall", "overall")
            table_rows.append(
                [
                    METHOD_NAME[method],
                    pct(r["evidence_recall_at_k"]),
                    pct(r["complete_recall_at_k"]),
                    pct(r["mrr"]),
                    f'{float(r["tool_calls"]):.2f}',
                    f'{float(r["latency_ms"]):.2f}',
                    f'{float(r["latency_ms_p95"]):.2f}',
                ]
            )
        add_table(doc, ["方法", f"Evidence@{k}", f"Complete@{k}", "MRR", "Avg Calls", "Avg Latency ms", "P95 ms"], table_rows)

    agent_c5 = metric(rows, "two_stage_agent", 5, "overall", "overall", "complete_recall_at_k")
    fusion_c5 = metric(rows, "fusion", 5, "overall", "overall", "complete_recall_at_k")
    dense_c5 = metric(rows, "dense", 5, "overall", "overall", "complete_recall_at_k")
    file_c5 = metric(rows, "file_fts", 5, "overall", "overall", "complete_recall_at_k")
    graph_c5 = metric(rows, "graph_path", 5, "overall", "overall", "complete_recall_at_k")
    add_para(
        doc,
        f"总体上，two_stage_agent 的 Complete@5 为 {pct_value(agent_c5)}，高于 dense 的 {pct_value(dense_c5)}、file_fts 的 {pct_value(file_c5)}、graph_path 的 {pct_value(graph_c5)}，并与 fusion 的 {pct_value(fusion_c5)} 持平。不同于 fusion 固定调用三个核心工具，two_stage_agent 平均只调用 1.44 个工具，因此在完整证据召回和调用成本之间取得了更好的平衡。",
    )

    add_heading(doc, "8. 三类任务表现", 1)
    for k in [3, 5]:
        add_para(doc, f"三类任务 E/C@{k}，单元格为 Evidence / Complete：")
        table_rows = []
        for method in METHOD_ORDER:
            table_rows.append(
                [
                    METHOD_NAME[method],
                    ec_cell(rows, method, k, "semantic_fact"),
                    ec_cell(rows, method, k, "multi_hop_relation"),
                    ec_cell(rows, method, k, "exact_file_lookup"),
                    f'{metric(rows, method, k, "overall", "overall", "tool_calls"):.2f}',
                ]
            )
        add_table(doc, ["方法", "semantic_fact", "multi_hop_relation", "exact_file_lookup", "Avg Calls"], table_rows)

    add_para(
        doc,
        "从任务分布看，file_fts 在 semantic_fact 与 exact_file_lookup 上非常强，尤其 exact_file_lookup@5 达到 96.67% / 96.67%。但在 multi_hop_relation 上，file_fts 的 Evidence@5 为 100.00%，Complete@5 只有 41.67%，说明它经常能找到一个相关证据，但难以补齐多跳所需的全部文档。",
    )
    add_para(
        doc,
        "Graph Path 在 multi_hop_relation 上明显更适合，Complete@5 为 63.33%，高于 dense 和 file_fts；但它在 semantic_fact 与 exact_file_lookup 上不如 file_fts 稳定，说明纯图方法并不能统一覆盖所有任务类型。",
    )
    add_para(
        doc,
        "two_stage_agent 的优势在于按需组合：semantic_fact 和 exact_file_lookup 维持 96.67% / 96.67%，multi_hop_relation 的 Complete@5 达到 63.33%，与 graph_path 和 fusion 持平。它没有依赖固定全量融合，而是通过 Evidence Judge 判断何时补 graph、何时补 file、何时停止。",
    )

    add_heading(doc, "9. 与基线的全面对比", 1)
    add_table(
        doc,
        ["比较对象", "主要差异", "实验结论"],
        [
            ["Dense vs Agent", "Dense 语义召回强，但 multi-hop Complete@5 仅 45.00%。Agent 可在需要时补图检索。", "Agent Complete@5 85.56%，明显高于 Dense 的 74.44%。"],
            ["File FTS vs Agent", "File FTS 在文件与精确定位上强，但 multi-hop Complete@5 仅 41.67%。", "Agent 保留文件任务稳定性，同时把 multi-hop Complete@5 提升到 63.33%。"],
            ["Graph Path vs Agent", "Graph Path 适合多跳，但事实题和文件题不如 file/dense 稳。", "Agent 在总体 Complete@5 上高于 Graph Path，且三类任务更均衡。"],
            ["Rule Router vs Agent", "Rule Router 只做一次工具选择，不会在证据不足时补检索。", "Agent 通过 Evidence Judge 和二轮补检索，把总体 Complete@5 从 82.22% 提到 85.56%。"],
            ["Fusion vs Agent", "Fusion 固定调用三种工具，召回强但成本高。", "Agent 与 Fusion Complete@5 持平，Avg Calls 为 1.44，而 Fusion 为 3.00。"],
            ["Oracle vs Agent", "Oracle 使用 gold 选择最佳结果，仅作理论上界。", "Oracle Complete@5 为 90.00%，说明仍存在可改进空间。"],
        ],
    )

    add_heading(doc, "10. 成本与效率", 1)
    table_rows = []
    for method in METHOD_ORDER:
        r = find_row(rows, method, 5, "overall", "overall")
        table_rows.append(
            [
                METHOD_NAME[method],
                f'{float(r["tool_calls"]):.2f}',
                f'{float(r["latency_ms"]):.2f}',
                f'{float(r["latency_ms_p95"]):.2f}',
                pct(r["complete_recall_at_k"]),
                pct(r["mrr"]),
            ]
        )
    add_table(doc, ["方法", "Avg Calls", "Avg Latency ms", "P95 ms", "Complete@5", "MRR"], table_rows)
    add_para(
        doc,
        "从调用次数看，two_stage_agent 平均调用 1.44 次核心检索工具，明显低于 fusion 的 3.00 次。Latency 数值受远程 dense embedding 调用影响较大，因此报告中更推荐把 Average Tool Calls 作为方法成本主指标，把 Latency 作为工程实现参考。后续如果使用本地 embedding、缓存 query embedding 或批处理请求，时延可进一步下降。",
    )

    add_heading(doc, "11. 实现文件说明", 1)
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["code/src/router.py", "Rule Router：问题信号提取、file/graph/dense 分数、首轮工具选择。"],
            ["code/src/evidence_judge.py", "Evidence Judge：证据覆盖判断、停/继续决策、二轮工具选择、查询改写。"],
            ["code/src/agent.py", "Two-Stage Agent：路由、首轮检索、Judge、二轮检索、protected weighted RRF、预测输出。"],
            ["code/src/file_fts_search.py", "File FTS 检索器：FTS5、chunk 聚合、标题/短语/数字日期增强。"],
            ["code/src/graph_path_search.py", "Graph Path 检索器：实体文档图、路径 beam search、路径评分。"],
            ["code/src/dense_search.py", "Dense 检索器：MiniLM embedding、向量相似度检索。"],
            ["code/src/experiment_box.py", "反污染实验盒：公开运行视图、预测写出、gold-only 评价边界。"],
            ["results/final/", "冻结预测、逐题指标、聚合 summary 和 main_results.csv。"],
            ["experiment_records/", "汇报用 CSV 表、方法清单、运行命令记录。"],
            ["manifest/freeze_manifest.json", "冻结文件哈希、数据规模、方法列表。"],
        ],
    )

    add_heading(doc, "12. 系统边界与不足", 1)
    add_para(doc, "第一，multi_hop_relation 仍是最难任务。two_stage_agent 的 multi-hop Complete@5 为 63.33%，Complete@3 为 46.67%，说明系统能补充多跳证据，但 Top-3 内完整证据组合仍不够稳定。")
    add_para(doc, "第二，exact_file_lookup 上存在一定过度二轮调用。该任务 Complete@5 很高，但 Avg Calls 为 1.65，说明 Evidence Judge 对部分文件题过于谨慎。若未来追求更低成本，可在验证集上增强 exact/file anchor 的停止规则，但不应在 final 上继续调参。")
    add_para(doc, "第三，图检索依赖实体抽取和图构建质量。若实体抽取漏掉关键别名、简称或跨句关系，Graph Path 和 Agent 的多跳能力会受影响。")
    add_para(doc, "第四，Oracle 仍高于 Agent。Oracle 的 Complete@5 为 90.00%，Agent 为 85.56%，说明理论上还存在改进空间，但 Oracle 使用 gold 信息，不是可部署系统。")
    add_para(doc, "第五，时延目前受 DeepInfra 远程 embedding 请求影响较大。该问题属于工程实现成本，不代表 Agent 逻辑必须很慢；缓存和本地化部署可显著改善。")

    add_heading(doc, "13. 最终结论", 1)
    add_para(
        doc,
        "最终 two_stage_agent 证明了轻量级任务感知多索引智能体搜索的核心价值：它不固定调用所有检索器，而是先用 Rule Router 选择最可能有效的索引，再由 Evidence Judge 判断是否需要补检索，最后通过 protected weighted RRF 平衡首轮高置信证据和二轮补充证据。",
    )
    add_para(
        doc,
        "实验上，two_stage_agent 在 Complete@5 上达到 85.56%，与全量 fusion 持平，同时平均工具调用只有 1.44 次；相比 dense、file_fts、graph_path 三个单一基线，它在总体完整证据召回上更强，并且能在 semantic_fact、multi_hop_relation、exact_file_lookup 三类任务之间取得更均衡表现。",
    )
    add_para(
        doc,
        "因此，该系统适合作为最终方法进入报告：它的重点不是在单一指标上压倒所有基线，而是在保持较好证据召回的同时显著控制检索成本，并通过智能体式证据判断弥补不同检索器的结构性短板。",
    )

    return doc


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 18, True)


def add_center(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 15 if level == 1 else 12.5, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, 8.8)
    doc.add_paragraph()


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def find_row(rows: list[dict[str, str]], method: str, k: int, group_type: str, group_name: str) -> dict[str, str]:
    for row in rows:
        if row["method"] == method and int(row["k"]) == k and row["group_type"] == group_type and row["group_name"] == group_name:
            return row
    raise KeyError((method, k, group_type, group_name))


def metric(rows: list[dict[str, str]], method: str, k: int, group_type: str, group_name: str, key: str) -> float:
    return float(find_row(rows, method, k, group_type, group_name)[key])


def pct(value: str) -> str:
    return pct_value(float(value))


def pct_value(value: float) -> str:
    return f"{value * 100:.2f}%"


def ec_cell(rows: list[dict[str, str]], method: str, k: int, task: str) -> str:
    row = find_row(rows, method, k, "task_type", task)
    return f'{pct(row["evidence_recall_at_k"])} / {pct(row["complete_recall_at_k"])}'


if __name__ == "__main__":
    main()
